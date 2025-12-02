from docx import Document

# ==============================
# 1. 请假制度
# ==============================
doc1 = Document()
doc1.add_heading('星渊科技内部请假制度', level=1)

doc1.add_heading('一、总则', level=2)
doc1.add_paragraph("1. 本制度适用于星渊科技所有全职员工。")
doc1.add_paragraph("2. 员工请假需提前按照流程提交申请，并获得直属主管批准后方可生效。")
doc1.add_paragraph("3. 无故缺勤将按照旷工处理。")

doc1.add_heading('二、请假类型与说明', level=2)
doc1.add_paragraph("【事假】")
doc1.add_paragraph("1. 事假需提前至少 1 天提交。")
doc1.add_paragraph("2. 事假不带薪。")
doc1.add_paragraph("3. 累计事假超过 5 天/月需额外说明理由。")

doc1.add_paragraph("【病假】")
doc1.add_paragraph("1. 病假超过 1 天需医院证明。")
doc1.add_paragraph("2. 每年 5 天带薪病假，超出部分半薪。")

doc1.add_paragraph("【年假期间请假】")
doc1.add_paragraph("1. 突发情况需中断年假，应及时通知主管。")
doc1.add_paragraph("2. 未报备者视为事假。")

doc1.add_paragraph("【调休假】")
doc1.add_paragraph("1. 加班可折算调休。")
doc1.add_paragraph("2. 调休需 90 天内使用。")

doc1.add_paragraph("【婚假、产假、陪产假】")
doc1.add_paragraph("1. 婚假 3 日。")
doc1.add_paragraph("2. 产假不少于 98 天。")
doc1.add_paragraph("3. 陪产假 7 日。")

doc1.add_heading('三、请假流程', level=2)
doc1.add_paragraph("1. 系统提交申请。")
doc1.add_paragraph("2. 主管 24 小时内审批。")
doc1.add_paragraph("3. 涉及跨部门需抄送。")
doc1.add_paragraph("4. 人力负责留档。")

doc1.add_heading('四、例外条款', level=2)
doc1.add_paragraph("1. 紧急情况下可事后 24 小时内补交。")
doc1.add_paragraph("2. 伪造证明将按纪律处分。")

doc1.save("leave_policy.docx")


# ==============================
# 2. 年假制度
# ==============================
doc2 = Document()
doc2.add_heading('星渊科技员工年假管理制度', level=1)

doc2.add_heading('一、年假天数计算', level=2)
doc2.add_paragraph("1. 入职未满一年无当年年假。")
doc2.add_paragraph("2. 1-10 年：5 天。")
doc2.add_paragraph("3. 10 年以上：10 天。")
doc2.add_paragraph("4. 年假不跨年累计。")

doc2.add_heading('二、申请规则', level=2)
doc2.add_paragraph("1. 提前 3 天申请。")
doc2.add_paragraph("2. 项目高峰期主管可调整时间。")
doc2.add_paragraph("3. 不得拆分为不足 0.5 天。")
doc2.add_paragraph("4. 超过 5 天需负责人批准。")

doc2.add_heading('三、未休处理', level=2)
doc2.add_paragraph("1. 因工作未休可折算薪资（200%）。")
doc2.add_paragraph("2. 个人原因未休不补偿。")
doc2.add_paragraph("3. 不得跨年。")

doc2.add_heading('四、离职规定', level=2)
doc2.add_paragraph("1. 多休需工资扣回。")
doc2.add_paragraph("2. 未休按政策补偿。")

doc2.save("annual_leave.docx")


# ==============================
# 3. 加班与调休制度
# ==============================
doc3 = Document()
doc3.add_heading('星渊科技加班与调休管理制度', level=1)

doc3.add_heading('一、加班类型', level=2)
doc3.add_paragraph("1. 工作日加班需主管确认。")
doc3.add_paragraph("2. 休息日加班优先调休。")
doc3.add_paragraph("3. 法定节假日加班按 300% 支付。")

doc3.add_heading('二、调休规则', level=2)
doc3.add_paragraph("1. 工作日加班 2 小时以下折算 1 小时调休。")
doc3.add_paragraph("2. 超过 2 小时按实际折算。")
doc3.add_paragraph("3. 休息日 1:1 调休。")
doc3.add_paragraph("4. 90 天内使用，过期作废。")

doc3.add_heading('三、申请流程', level=2)
doc3.add_paragraph("1. 提前提交加班申请。")
doc3.add_paragraph("2. 调休需标注加班编号。")
doc3.add_paragraph("3. 人力更新调休余额。")

doc3.add_heading('四、例外情况', level=2)
doc3.add_paragraph("1. 特殊情况可申请延长 30 天。")
doc3.add_paragraph("2. 未批准加班不计调休。")

doc3.save("overtime_policy.docx")

print("已成功生成三个 docx 文件！")