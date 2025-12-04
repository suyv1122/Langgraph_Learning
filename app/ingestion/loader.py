from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
import docx
# from app.config import settings
from app.config import settings

def load_pdf(path: Path) -> List[Document]:
    """下面的代码将pdf分割成单独的页面，每一个页面的文本被封装成一个Document放入list"""
    reader = PdfReader(str(path))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(path), "page": i+1}
            ))
    return docs

def load_docx(path: Path) -> List[Document]:
    """
    读取指定目录下所有支持的文档。
    如果 dir_path 是相对路径，则自动从项目根目录拼接。
    如不传路径，默认使用 settings.docs_dir
    """
    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": str(path)})] if text else []



def load_docs(dir_path: str | Path | None = None) -> List[Document]:
    # p = Path(dir_path)
    # docs: List[Document] = []
    # for f in p.rglob("*"):
    #     if f.suffix.lower() == ".pdf":
    #         docs.extend(load_pdf(f))
    #     elif f.suffix.lower() in [".docx", ".doc"]:
    #         docs.extend(load_docx(f))
    #     elif f.suffix.lower() in [".md", ".txt"]:
    #         docs.append(Document(page_content=f.read_text(encoding="utf-8"),
    #                              metadata={"source": str(f)}))
    # return docs

    # 决定原始路径
    if dir_path is None:
        p = Path(settings.docs_dir)
    else:
        p = Path(dir_path)

    # 若是相对路径，均视为自项目根目录开始的相对路径
    if not p.is_absolute():
        # 自当前文件向上两级，去项目根目录
        project_root = Path(__file__).resolve().parents[2]
        p = project_root / p

    # 扫描文件
    docs: List[Document] = []
    for f in p.rglob("*"):
        if f.suffix.lower() == ".pdf":
            docs.extend(load_pdf(f))
        elif f.suffix.lower() in [".docx", ".doc"]:
            docs.extend(load_docx(f))
        elif f.suffix.lower() in [".md", ".txt"]:
            docs.append(Document(
                page_content=f.read_text(encoding="utf-8"),
                metadata={"source": str(f)}
            ))
    return docs

def split_docs(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap
    )
    return splitter.split_documents(docs)

def load_single_file(path: Path) -> List[Document]:
    """用于应对后续文件单独追加而设计，本身用于处理单个文件的读取"""
    """根据文件后缀加载文件，返回LangChain的Document列表"""
    suf = path.suffix.lower()
    if suf == '.pdf':
        return load_pdf(path)
    if suf in ['.docx', '.doc']:
        return load_docx(path)
    if suf in ['.md', '.txt']:
        text = path.read_text(encoding="utf-8")
        return [Document(page_content=text, metadata={"source": str(path)})] if text.strip() else []
    return []

def split_with_visibility(docs: List[Document], visibility: str, doc_id: str | None = None) -> List[Document]:
    chunks = split_docs(docs)
    for c in chunks:
        c.metadata = dict(c.metadata or {})
        c.metadata["visibility"] = visibility
        if doc_id:
            c.metadata["doc_id"] = doc_id
    return chunks